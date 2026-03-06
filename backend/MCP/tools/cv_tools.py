"""
CV tools for HR Agent.
Provides employee CV retrieval, LLM-powered analysis & summary,
and file management (upload/replace/delete).
"""
import os
import json
import re
import shutil
import traceback
from datetime import datetime
from typing import Dict, Any, Optional
from pathlib import Path
import cx_Oracle
from dotenv import load_dotenv
import ollama

load_dotenv()

ORACLE_USER = os.getenv("ORACLE_USER")
ORACLE_PASSWORD = os.getenv("ORACLE_PASSWORD")
ORACLE_HOST = os.getenv("ORACLE_HOST")
ORACLE_PORT = os.getenv("ORACLE_PORT")
ORACLE_SERVICE = os.getenv("ORACLE_SERVICE_NAME")

dsn = cx_Oracle.makedsn(ORACLE_HOST, ORACLE_PORT, service_name=ORACLE_SERVICE)

# LLM Model for analysis
ANALYSIS_MODEL = os.getenv("ANALYSIS_MODEL", "granite4:350m")

# CV file storage — import from centralized config
import sys
sys.path.insert(0, str(Path(__file__).resolve().parent.parent.parent))
from config import CV_DIR, url_to_abs_path


def _get_connection():
    """Get Oracle database connection."""
    return cx_Oracle.connect(user=ORACLE_USER, password=ORACLE_PASSWORD, dsn=dsn)


def _read_lob(val):
    """Read Oracle LOB value if needed."""
    if val and hasattr(val, 'read'):
        return val.read()
    return val


def _format_currency(amount) -> str:
    """Format number as Indonesian Rupiah."""
    if amount is None:
        return "Rp 0"
    return f"Rp {int(amount):,}".replace(",", ".")


def _get_cv_data(cur, emp_id: int) -> Optional[Dict]:
    """Helper to fetch and format CV data from DB."""
    cur.execute("""
        SELECT cv.*, e.name, e.employee_code, e.email, e.phone,
               e.marital_status, e.employment_status, e.joined_at, e.address
        FROM employee_cv cv
        JOIN employees e ON cv.employee_id = e.id
        WHERE cv.employee_id = :eid
    """, {"eid": emp_id})
    
    columns = [desc[0].lower() for desc in cur.description]
    row = cur.fetchone()
    
    if not row:
        return None
    
    data = {}
    for i, col in enumerate(columns):
        val = row[i]
        data[col] = _read_lob(val) if val else val
    
    return data


# ==================== TOOL 1: GET EMPLOYEE CV ====================

def get_employee_cv(emp_id: int) -> Dict[str, Any]:
    """
    Ambil data CV/profil lengkap karyawan dari database.
    
    Args:
        emp_id: Database ID karyawan
        
    Returns:
        Dict with complete CV data
    """
    try:
        conn = _get_connection()
        cur = conn.cursor()
        
        # Check employee exists
        cur.execute("SELECT name FROM employees WHERE id = :eid", {"eid": emp_id})
        emp = cur.fetchone()
        if not emp:
            cur.close(); conn.close()
            return {"success": False, "error": f"Karyawan dengan ID {emp_id} tidak ditemukan."}
        
        cv_data = _get_cv_data(cur, emp_id)
        cur.close(); conn.close()
        
        if not cv_data:
            return {"success": False, "error": f"CV karyawan {emp[0]} belum tersedia di database."}
        
        # Format response
        result = {
            "success": True,
            "employee": {
                "id": emp_id,
                "name": cv_data.get("name"),
                "code": cv_data.get("employee_code"),
                "email": cv_data.get("email"),
                "phone": cv_data.get("phone"),
                "marital_status": cv_data.get("marital_status"),
                "employment_status": cv_data.get("employment_status"),
                "joined_at": cv_data.get("joined_at").strftime("%Y-%m-%d") if cv_data.get("joined_at") and hasattr(cv_data.get("joined_at"), 'strftime') else str(cv_data.get("joined_at", "")),
                "address": cv_data.get("address"),
            },
            "current_info": {
                "position": cv_data.get("current_position"),
                "department": cv_data.get("current_department"),
                "salary": _format_currency(cv_data.get("current_salary")),
            },
            "education": {
                "level": cv_data.get("education_level"),
                "institution": cv_data.get("education_institution"),
                "major": cv_data.get("education_major"),
                "graduation_year": cv_data.get("graduation_year"),
            },
            "qualifications": {
                "certifications": cv_data.get("certifications"),
                "skills": cv_data.get("skills"),
                "work_experience": cv_data.get("work_experience"),
            },
            "personal_data": {
                "blood_type": cv_data.get("blood_type"),
                "religion": cv_data.get("religion"),
                "ktp_number": cv_data.get("ktp_number"),
                "npwp_number": cv_data.get("npwp_number"),
            },
            "emergency_contact": {
                "name": cv_data.get("emergency_contact_name"),
                "phone": cv_data.get("emergency_contact_phone"),
                "relation": cv_data.get("emergency_contact_relation"),
            },
            "bank_info": {
                "bank_name": cv_data.get("bank_name"),
                "account_number": cv_data.get("bank_account_number"),
                "account_name": cv_data.get("bank_account_name"),
            },
            "monthly_deductions": {
                "bpjs_kesehatan": _format_currency(cv_data.get("deduction_bpjs_kesehatan")),
                "bpjs_ketenagakerjaan": _format_currency(cv_data.get("deduction_bpjs_ketenagakerjaan")),
                "meal": _format_currency(cv_data.get("deduction_meal")),
                "transport": _format_currency(cv_data.get("deduction_transport")),
                "insurance": _format_currency(cv_data.get("deduction_insurance")),
                "laptop_installment": _format_currency(cv_data.get("deduction_laptop_installment")),
                "laptop_remaining_months": cv_data.get("deduction_laptop_remaining_months"),
                "other": _format_currency(cv_data.get("deduction_other")),
                "other_description": cv_data.get("deduction_other_description"),
                "total_monthly": _format_currency(cv_data.get("total_monthly_deductions")),
            },
            "cv_file": {
                "file_path": cv_data.get("file_path"),
                "has_file": bool(cv_data.get("file_path") and Path(str(cv_data.get("file_path"))).exists()) if cv_data.get("file_path") else False
            }
        }
        
        return result
        
    except Exception as e:
        return {"success": False, "error": str(e), "trace": traceback.format_exc()}


# ==================== TOOL 2: ANALYZE EMPLOYEE CV (LLM) ====================

def analyze_employee_cv(emp_id: int, focus: str = "general") -> Dict[str, Any]:
    """
    Analisa mendalam CV karyawan menggunakan AI.
    
    Args:
        emp_id: Database ID karyawan
        focus: Fokus analisa: "general", "skills", "career", "compensation", "performance"
        
    Returns:
        Dict with AI analysis results
    """
    try:
        conn = _get_connection()
        cur = conn.cursor()
        
        cv_data = _get_cv_data(cur, emp_id)
        cur.close(); conn.close()
        
        if not cv_data:
            return {"success": False, "error": f"CV karyawan dengan ID {emp_id} tidak ditemukan."}
        
        # Build CV text for LLM
        cv_text = f"""
PROFIL KARYAWAN:
- Nama: {cv_data.get('name')}
- Kode: {cv_data.get('employee_code')}
- Posisi: {cv_data.get('current_position')} di {cv_data.get('current_department')}
- Status: {cv_data.get('employment_status')}
- Gaji: Rp {cv_data.get('current_salary', 0):,}
- Status Pernikahan: {cv_data.get('marital_status')}

PENDIDIKAN:
- {cv_data.get('education_level')} - {cv_data.get('education_institution')}
- Jurusan: {cv_data.get('education_major')} (Lulus: {cv_data.get('graduation_year')})

SERTIFIKASI:
{cv_data.get('certifications', 'Tidak ada')}

KEAHLIAN:
{cv_data.get('skills', 'Tidak ada')}

PENGALAMAN KERJA:
{cv_data.get('work_experience', 'Tidak ada')}

POTONGAN BULANAN:
- BPJS Kes: Rp {cv_data.get('deduction_bpjs_kesehatan', 0):,}
- BPJS TK: Rp {cv_data.get('deduction_bpjs_ketenagakerjaan', 0):,}
- Asuransi: Rp {cv_data.get('deduction_insurance', 0):,}
- Cicilan Laptop: Rp {cv_data.get('deduction_laptop_installment', 0):,} ({cv_data.get('deduction_laptop_remaining_months', 0)} bulan tersisa)
- Total Potongan: Rp {cv_data.get('total_monthly_deductions', 0):,}
"""
        
        # Read CV file content if exists
        file_content = ""
        file_path = cv_data.get("file_path")
        if file_path and Path(str(file_path)).exists():
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    file_content = f.read()[:5000]  # Limit to 5000 chars
                cv_text += f"\n\nISI FILE CV:\n{file_content}"
            except Exception:
                pass  # Skip if file can't be read as text
        
        focus_instructions = {
            "general": "Berikan analisa menyeluruh tentang profil karyawan ini.",
            "skills": "Fokus pada analisa keahlian, sertifikasi, dan gap kompetensi.",
            "career": "Fokus pada jalur karir, pengalaman, dan potensi pengembangan.",
            "compensation": "Fokus pada analisa kompensasi versus pengalaman dan posisi.",
            "performance": "Fokus pada potensi kinerja berdasarkan profil."
        }
        
        analysis_prompt = f"""Analisa CV karyawan berikut:

{cv_text}

## Instruksi Analisis
{focus_instructions.get(focus, focus_instructions['general'])}

Berikan analisa dalam format JSON:
{{
  "summary": "ringkasan profil karyawan",
  "strengths": ["kekuatan 1", "kekuatan 2"],
  "weaknesses": ["kelemahan/gap 1", "kelemahan/gap 2"],
  "recommendations": ["rekomendasi 1", "rekomendasi 2"],
  "career_potential": "penilaian potensi karir",
  "skill_alignment": "seberapa sesuai keahlian dengan posisi saat ini",
  "compensation_assessment": "penilaian kompensasi"
}}"""

        response = ollama.generate(
            model=ANALYSIS_MODEL,
            prompt=analysis_prompt,
            options={"temperature": 0.3}
        )
        
        analysis_text = response.get("response", "")
        
        # Parse JSON
        json_match = re.search(r'\{[\s\S]*\}', analysis_text)
        if json_match:
            try:
                analysis_result = json.loads(json_match.group())
            except json.JSONDecodeError:
                analysis_result = {"summary": analysis_text, "parse_error": True}
        else:
            analysis_result = {"summary": analysis_text, "parse_error": True}
        
        return {
            "success": True,
            "employee": cv_data.get("name"),
            "focus": focus,
            "analysis": analysis_result,
            "has_cv_file": bool(file_content)
        }
        
    except Exception as e:
        return {"success": False, "error": str(e), "trace": traceback.format_exc()}


# ==================== TOOL 3: SUMMARIZE/ASK ABOUT CV (LLM) ====================

def summarize_employee_cv(emp_id: int, question: Optional[str] = None) -> Dict[str, Any]:
    """
    Rangkum CV karyawan atau jawab pertanyaan spesifik tentang CV.
    
    Args:
        emp_id: Database ID karyawan
        question: Pertanyaan spesifik tentang CV (opsional, jika None = rangkuman umum)
        
    Returns:
        Dict with summary or answer
    """
    try:
        conn = _get_connection()
        cur = conn.cursor()
        
        cv_data = _get_cv_data(cur, emp_id)
        cur.close(); conn.close()
        
        if not cv_data:
            return {"success": False, "error": f"CV karyawan dengan ID {emp_id} tidak ditemukan."}
        
        # Build CV text
        cv_text = f"""
NAMA: {cv_data.get('name')}
KODE: {cv_data.get('employee_code')}
POSISI: {cv_data.get('current_position')} - {cv_data.get('current_department')}
STATUS KERJA: {cv_data.get('employment_status')}
GAJI: Rp {cv_data.get('current_salary', 0):,}
BERGABUNG: {cv_data.get('joined_at')}

PENDIDIKAN: {cv_data.get('education_level')} - {cv_data.get('education_institution')} - {cv_data.get('education_major')} ({cv_data.get('graduation_year')})
SERTIFIKASI: {cv_data.get('certifications', 'Tidak ada')}
KEAHLIAN: {cv_data.get('skills', 'Tidak ada')}
PENGALAMAN: {cv_data.get('work_experience', 'Tidak ada')}

DATA PRIBADI:
- KTP: {cv_data.get('ktp_number')}
- NPWP: {cv_data.get('npwp_number')}
- Golongan Darah: {cv_data.get('blood_type')}
- Agama: {cv_data.get('religion')}
- Status: {cv_data.get('marital_status')}

KONTAK DARURAT: {cv_data.get('emergency_contact_name')} ({cv_data.get('emergency_contact_relation')}) - {cv_data.get('emergency_contact_phone')}

BANK: {cv_data.get('bank_name')} - {cv_data.get('bank_account_number')} a.n. {cv_data.get('bank_account_name')}

POTONGAN BULANAN:
- BPJS Kesehatan: Rp {cv_data.get('deduction_bpjs_kesehatan', 0):,}
- BPJS TK: Rp {cv_data.get('deduction_bpjs_ketenagakerjaan', 0):,}
- Makan: Rp {cv_data.get('deduction_meal', 0):,}
- Transport: Rp {cv_data.get('deduction_transport', 0):,}
- Asuransi: Rp {cv_data.get('deduction_insurance', 0):,}
- Cicilan Laptop: Rp {cv_data.get('deduction_laptop_installment', 0):,} ({cv_data.get('deduction_laptop_remaining_months', 0)} bln)
- Lainnya: Rp {cv_data.get('deduction_other', 0):,} ({cv_data.get('deduction_other_description', '')})
- Total: Rp {cv_data.get('total_monthly_deductions', 0):,}
"""
        
        # Read CV file content if exists
        file_path = cv_data.get("file_path")
        if file_path and Path(str(file_path)).exists():
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    file_content = f.read()[:5000]
                cv_text += f"\nISI FILE CV:\n{file_content}"
            except Exception:
                pass
        
        if question:
            prompt = f"""Berdasarkan CV karyawan berikut, jawab pertanyaan ini:

{cv_text}

PERTANYAAN: {question}

Jawab dengan lengkap dan informatif berdasarkan data CV di atas. Jika informasi tidak tersedia, katakan demikian."""
        else:
            prompt = f"""Rangkum CV karyawan berikut secara komprehensif:

{cv_text}

Buat rangkuman yang mencakup:
1. Identitas dan posisi saat ini
2. Latar belakang pendidikan dan kualifikasi
3. Pengalaman kerja dan keahlian utama
4. Informasi kompensasi dan potongan
5. Highlight atau catatan penting"""
        
        response = ollama.generate(
            model=ANALYSIS_MODEL,
            prompt=prompt,
            options={"temperature": 0.3}
        )
        
        result_text = response.get("response", "")
        
        return {
            "success": True,
            "employee": cv_data.get("name"),
            "type": "answer" if question else "summary",
            "question": question,
            "result": result_text,
            "has_cv_file": bool(file_path and Path(str(file_path)).exists())
        }
        
    except Exception as e:
        return {"success": False, "error": str(e), "trace": traceback.format_exc()}


# ==================== TOOL 4: MANAGE CV FILE ====================

def manage_cv_file(emp_id: int, action: str, file_path: Optional[str] = None) -> Dict[str, Any]:
    """
    Kelola file CV karyawan: upload, replace, atau delete.
    
    Args:
        emp_id: Database ID karyawan
        action: "upload", "replace", atau "delete"
        file_path: Path ke file yang akan di-upload/replace (wajib untuk upload/replace)
        
    Returns:
        Dict with success status and file info
    """
    try:
        conn = _get_connection()
        cur = conn.cursor()
        
        # Check employee exists
        cur.execute("SELECT name, employee_code FROM employees WHERE id = :eid", {"eid": emp_id})
        emp = cur.fetchone()
        if not emp:
            cur.close(); conn.close()
            return {"success": False, "error": f"Karyawan dengan ID {emp_id} tidak ditemukan."}
        
        emp_name, emp_code = emp
        
        # Get existing CV record
        cur.execute("SELECT id, file_path FROM employee_cv WHERE employee_id = :eid", {"eid": emp_id})
        cv_row = cur.fetchone()
        
        if action == "delete":
            if cv_row and cv_row[1]:
                old_path = Path(str(cv_row[1]))
                if old_path.exists():
                    os.remove(str(old_path))
                
                cur.execute("""
                    UPDATE employee_cv SET file_path = NULL, download_url = NULL, updated_at = CURRENT_TIMESTAMP
                    WHERE employee_id = :eid
                """, {"eid": emp_id})
                conn.commit()
                cur.close(); conn.close()
                
                return {"success": True, "message": f"File CV {emp_name} berhasil dihapus."}
            else:
                cur.close(); conn.close()
                return {"success": False, "error": f"Tidak ada file CV untuk {emp_name}."}
        
        elif action in ("upload", "replace"):
            if not file_path:
                cur.close(); conn.close()
                return {"success": False, "error": "Parameter file_path wajib untuk upload/replace."}
            
            # Clean up the path string just in case LLM added quotes
            file_path = file_path.strip('\'"')
            
            # Resolve URL ke absolute path jika input berupa URL (misal dari frontend)
            from config import url_to_abs_path, BASE_URL
            resolved_abs_path = url_to_abs_path(file_path)
            if resolved_abs_path:
                source_path = resolved_abs_path
            else:
                source_path = Path(file_path).resolve()
                
            if not source_path.exists() or not source_path.is_file():
                cur.close(); conn.close()
                return {
                    "success": False,
                    "error": f"File tidak ditemukan atau path adalah sebuah direktori: {file_path}",
                    "debug": {
                        "given_path": file_path,
                        "resolved_path": str(source_path),
                        "cwd": os.getcwd(),
                        "cv_dir": str(CV_DIR),
                        "hint": f"Pastikan file_path mengarah ke spesifik file, bukan hanya foldernya."
                    }
                }
            
            # Delete old file if replacing
            if action == "replace" and cv_row and cv_row[1]:
                old_path = Path(str(cv_row[1])).resolve()
                if old_path.exists():
                    os.remove(str(old_path))
            
            # Copy file to CV directory (skip copy if source == destination)
            safe_name = emp_name.replace(" ", "_")
            ext = source_path.suffix or ".pdf"
            dest_filename = f"CV_{safe_name}_{emp_code}{ext}"
            # Pastikan direktori tujuan ada
            CV_DIR.mkdir(parents=True, exist_ok=True)
            dest_path = (CV_DIR / dest_filename).resolve()
            
            if source_path != dest_path:
                shutil.copy2(str(source_path), str(dest_path))
            # else: already in the right place, no copy needed
            
            download_url = f"/api/uploads/cv/{dest_filename}"
            # PENTING: Simpan path absolut fisik di file_path, dan URL di download_url
            db_file_path = str(dest_path)
            
            if cv_row:
                cur.execute("""
                    UPDATE employee_cv SET file_path = :fp, download_url = :du, updated_at = CURRENT_TIMESTAMP
                    WHERE employee_id = :eid
                """, {"fp": db_file_path, "du": download_url, "eid": emp_id})
            else:
                # Create CV record if it doesn't exist
                cur.execute("SELECT name, department, position, basic_salary FROM employees WHERE id = :eid", {"eid": emp_id})
                emp_info = cur.fetchone()
                cur.execute("""
                    INSERT INTO employee_cv (employee_id, current_position, current_department, current_salary, file_path, download_url)
                    VALUES (:eid, :pos, :dept, :sal, :fp, :du)
                """, {
                    "eid": emp_id,
                    "pos": emp_info[2] if emp_info else None,
                    "dept": emp_info[1] if emp_info else None,
                    "sal": emp_info[3] if emp_info else 0,
                    "fp": db_file_path,
                    "du": download_url
                })
            
            conn.commit()
            cur.close(); conn.close()
            
            file_size = os.path.getsize(dest_path)
            
            return {
                "success": True,
                "message": f"File CV {emp_name} berhasil di-{action}.",
                "file_path": db_file_path,
                "download_url": download_url,
                "server_url": f"{BASE_URL}{download_url}",
                "filename": dest_filename,
                "widget": {
                    "type": "download",
                    "filename": dest_filename,
                    "size": f"{file_size / 1024:.1f} KB",
                    "icon": ext.replace(".", ""),
                    "download_url": download_url
                }
            }
        
        else:
            cur.close(); conn.close()
            return {"success": False, "error": f"Action tidak valid: {action}. Gunakan 'upload', 'replace', atau 'delete'."}
        
    except Exception as e:
        return {"success": False, "error": str(e), "trace": traceback.format_exc()}


# ==================== TOOL 5: EXTRACT CV DATA FROM FILE (LLM) ====================

def _read_file_content(file_path: str) -> str:
    """Read text content from PDF, DOCX, or TXT file."""
    path = Path(file_path)
    ext = path.suffix.lower()
    
    if ext == ".pdf":
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(str(path))
            text = ""
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            return text.strip()
        except Exception as e:
            # Fallback: try reading as plain text if PDF parsing fails
            # This helps with corrupted files or text files with .pdf extension
            try:
                with open(str(path), 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read().strip()
                    if content:
                        return content
                return f"[Error reading PDF: {e}]"
            except Exception:
                return f"[Error reading PDF: {e}]"
    
    elif ext in (".docx", ".doc"):
        try:
            from docx import Document
            doc = Document(str(path))
            text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            # Also read tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text += "\n" + row_text
            return text.strip()
        except Exception as e:
            return f"[Error reading DOCX: {e}]"
    
    elif ext in (".txt", ".text"):
        try:
            with open(str(path), 'r', encoding='utf-8', errors='ignore') as f:
                return f.read().strip()
        except Exception as e:
            return f"[Error reading TXT: {e}]"
    
    else:
        # Try reading as plain text
        try:
            with open(str(path), 'r', encoding='utf-8', errors='ignore') as f:
                return f.read().strip()
        except Exception:
            return ""


def extract_cv_from_file(emp_id: int, file_path: Optional[str] = None) -> Dict[str, Any]:
    """
    Baca file CV yang diupload, ekstrak data menggunakan AI, lalu simpan ke database.
    Hanya kolom yang ditemukan informasinya di CV yang akan diisi.
    
    Args:
        emp_id: Database ID karyawan
        file_path: Path file CV (opsional, jika None ambil dari record employee_cv)
        
    Returns:
        Dict with extracted fields and update status
    """
    try:
        conn = _get_connection()
        cur = conn.cursor()
        
        # Verify employee
        cur.execute("SELECT name, employee_code, department, position, basic_salary FROM employees WHERE id = :eid", {"eid": emp_id})
        emp = cur.fetchone()
        if not emp:
            cur.close(); conn.close()
            return {"success": False, "error": f"Karyawan dengan ID {emp_id} tidak ditemukan."}
        
        emp_name, emp_code, emp_dept, emp_pos, emp_salary = emp
        
        # Get file path from DB if not provided
        if not file_path:
            cur.execute("SELECT file_path FROM employee_cv WHERE employee_id = :eid", {"eid": emp_id})
            row = cur.fetchone()
            if row and row[0]:
                file_path = row[0]
            else:
                cur.close(); conn.close()
                return {"success": False, "error": f"Tidak ada file CV untuk {emp_name}. Upload file terlebih dahulu menggunakan manage_cv_file."}
        
        # Clean up the path string
        file_path = str(file_path).strip('\'"')

        # Resolve server URL to absolute path if needed
        if file_path and (file_path.startswith("http://") or file_path.startswith("https://")):
            from config import url_to_abs_path
            resolved = url_to_abs_path(file_path)
            if resolved:
                file_path = str(resolved)
            else:
                cur.close(); conn.close()
                return {"success": False, "error": f"Tidak bisa resolve server URL ke path: {file_path}"}

        # Verify file exists
        file_path_obj = Path(file_path).resolve()
        if not file_path_obj.exists() or not file_path_obj.is_file():
            cur.close(); conn.close()
            return {"success": False, "error": f"File tidak ditemukan atau path adalah direktori: {file_path}"}
        
        # Override file_path with absolute, resolved string
        file_path = str(file_path_obj)
        
        # Read file content
        file_content = _read_file_content(file_path)
        if not file_content or file_content.startswith("[Error"):
            cur.close(); conn.close()
            return {"success": False, "error": f"Gagal membaca isi file: {file_content}"}
        
        # Truncate if too long (20k chars is plenty for most CVs and faster for local LLM)
        if len(file_content) > 20000:
            file_content = file_content[:20000] + "\n... [dipotong karena terlalu panjang]"
        
        # Build LLM prompt
        extraction_prompt = f"""Kamu adalah AI asisten HR yang bertugas mengekstrak data dari CV/resume karyawan.

ISI CV:
\"\"\"
{file_content}
\"\"\"

INFORMASI KARYAWAN SAAT INI (dari sistem):
- Nama: {emp_name}
- Kode: {emp_code}
- Departemen: {emp_dept}
- Posisi: {emp_pos}

## Instruksi
Ekstrak HANYA informasi yang benar-benar ada/ditemukan di CV di atas. Jangan mengarang atau mengasumsikan data yang tidak ada.

Field yang bisa diekstrak (semua opsional, isi HANYA jika ada di CV):
- "education_level": tingkat pendidikan terakhir (contoh: "S1", "S2", "D3", "SMA")
- "education_institution": nama universitas/insitusi pendidikan
- "education_major": jurusan/program studi
- "graduation_year": tahun lulus (angka, contoh: 2020)
- "certifications": daftar sertifikasi/training (teks bebas, pisahkan dengan koma atau newline)
- "skills": daftar keahlian/skill teknis dan non-teknis
- "work_experience": ringkasan pengalaman kerja sebelumnya (kronologis, dengan jabatan dan durasi)
- "emergency_contact_name": nama kontak darurat
- "emergency_contact_phone": nomor telepon kontak darurat
- "emergency_contact_relation": hubungan (contoh: "Istri", "Suami", "Orang Tua")
- "blood_type": golongan darah (contoh: "A", "B", "AB", "O")
- "religion": agama
- "ktp_number": nomor KTP/NIK
- "npwp_number": nomor NPWP
- "bank_name": nama bank
- "bank_account_number": nomor rekening
- "bank_account_name": nama pemilik rekening
- "notes": catatan tambahan penting yang tidak masuk kategori di atas

## Format Output
Keluarkan HANYA JSON object. Jangan sertakan field yang datanya TIDAK ADA di CV.
Contoh output jika hanya menemukan pendidikan dan skill:
{{"education_level": "S1", "education_institution": "Universitas Indonesia", "education_major": "Teknik Informatika", "graduation_year": 2018, "skills": "Python, Java, SQL, Docker, Machine Learning"}}

PENTING: Output HANYA JSON, tanpa penjelasan tambahan."""

        try:
            response = ollama.generate(
                model=ANALYSIS_MODEL,
                prompt=extraction_prompt,
                options={"temperature": 0.1}  # Low temp for factual extraction
            )
        except Exception as oe:
            cur.close(); conn.close()
            error_msg = str(oe)
            if "not found" in error_msg.lower():
                error_msg = f"Model '{ANALYSIS_MODEL}' tidak ditemukan di Ollama. Silakan jalankan 'ollama pull {ANALYSIS_MODEL}'."
            return {"success": False, "error": f"Ollama generation failed: {error_msg}"}
        
        llm_output = response.get("response", "")
        
        # Parse JSON from LLM response
        json_match = re.search(r'\{[\s\S]*\}', llm_output)
        if not json_match:
            cur.close(); conn.close()
            return {"success": False, "error": "LLM tidak menghasilkan JSON yang valid.", "raw_output": llm_output}
        
        try:
            extracted = json.loads(json_match.group())
        except json.JSONDecodeError:
            cur.close(); conn.close()
            return {"success": False, "error": "Gagal parse JSON dari LLM.", "raw_output": llm_output}
        
        if not extracted:
            cur.close(); conn.close()
            return {"success": False, "error": "LLM tidak menemukan informasi yang dapat diekstrak dari CV."}
        
        # Whitelist of valid columns (only extractable fields, NO system/deduction fields)
        valid_columns = {
            "education_level", "education_institution", "education_major", "graduation_year",
            "certifications", "skills", "work_experience",
            "emergency_contact_name", "emergency_contact_phone", "emergency_contact_relation",
            "blood_type", "religion", "ktp_number", "npwp_number",
            "bank_name", "bank_account_number", "bank_account_name",
            "notes"
        }
        
        # Filter to only valid columns and ensure values are bindable (no lists for standard SQL)
        clean_data = {}
        for key, value in extracted.items():
            if key in valid_columns and value is not None:
                # Oracle ORA-01484 fix: join lists into strings
                if isinstance(value, list):
                    value = ", ".join(map(str, value))
                
                if str(value).strip():
                    clean_data[key] = value
        
        if not clean_data:
            cur.close(); conn.close()
            return {"success": False, "error": "Tidak ada field yang valid diekstrak dari CV."}
        
        # Check if CV record exists
        cur.execute("SELECT id FROM employee_cv WHERE employee_id = :eid", {"eid": emp_id})
        cv_exists = cur.fetchone()
        
        if cv_exists:
            # Build dynamic UPDATE
            set_clauses = []
            params = {"eid": emp_id}
            for key, value in clean_data.items():
                set_clauses.append(f"{key} = :{key}")
                params[key] = value
            
            set_clauses.append("updated_at = CURRENT_TIMESTAMP")
            
            update_sql = f"UPDATE employee_cv SET {', '.join(set_clauses)} WHERE employee_id = :eid"
            cur.execute(update_sql, params)
        else:
            # INSERT new record
            clean_data["employee_id"] = emp_id
            clean_data["current_position"] = emp_pos
            clean_data["current_department"] = emp_dept
            clean_data["current_salary"] = emp_salary or 0
            
            cols = ", ".join(clean_data.keys())
            placeholders = ", ".join([f":{k}" for k in clean_data.keys()])
            insert_sql = f"INSERT INTO employee_cv ({cols}) VALUES ({placeholders})"
            cur.execute(insert_sql, clean_data)
        
        conn.commit()
        cur.close()
        conn.close()
        
        return {
            "success": True,
            "message": f"Data CV {emp_name} berhasil diekstrak dan disimpan. {len(clean_data)} field diisi.",
            "data": clean_data, # For tool chaining (Stage 2 expects .data)
            "employee": {"id": emp_id, "name": emp_name, "code": emp_code},
            "extracted_fields": list(clean_data.keys()),
            "extracted_data": {k: (v[:100] + "..." if isinstance(v, str) and len(v) > 100 else v) for k, v in clean_data.items()},
            "source_file": Path(file_path).name,
            "mode": "update" if cv_exists else "insert"
        }
        
    except Exception as e:
        return {"success": False, "error": str(e), "trace": traceback.format_exc()}


# ==================== TOOL 6: UPDATE EMPLOYEE CV ====================

def update_employee_cv(emp_id: int, updates: Optional[Dict[str, Any]] = None, **kwargs) -> Dict[str, Any]:
    """
    Update employee CV details by ID.
    Supports only 'employee_cv' table.
    
    Args:
        emp_id: Employee ID
        updates: Dict of field names and values to update (optional)
        **kwargs: The columns to update, directly as keyword arguments
        
    Valid employee_cv fields: education_level, education_institution, education_major,
                              graduation_year, certifications, skills, work_experience,
                              current_position, current_department, current_salary,
                              emergency_contact_name, emergency_contact_phone,
                              emergency_contact_relation, blood_type, religion,
                              ktp_number, npwp_number, bank_name, bank_account_number,
                              bank_account_name, deduction_bpjs_kesehatan,
                              deduction_bpjs_ketenagakerjaan, deduction_meal,
                              deduction_transport, deduction_insurance,
                              deduction_laptop_installment, deduction_laptop_remaining_months,
                              deduction_other, deduction_other_description,
                              total_monthly_deductions, notes        
    Returns:
        Dict with success status and message
    """
    actual_updates = updates or {}
    actual_updates.update({k: v for k, v in kwargs.items() if v is not None})
    
    # Guard: must be a non-empty dict
    if not actual_updates:
        return {
            "success": False,
            "error": "Parameter update tidak boleh kosong. Pastikan field dan nilai yang ingin diupdate sudah ditentukan."
        }
    
    # Oracle ORA-01484 fix: Ensure no lists/arrays are passed as bound values to standard SQL
    for k, v in actual_updates.items():
        if isinstance(v, list):
            actual_updates[k] = ", ".join(map(str, v))

    try:
        conn = _get_connection()
        cur = conn.cursor()
        
        # Get valid columns for employee_cv table (dynamically)
        cur.execute("SELECT * FROM employee_cv WHERE ROWNUM = 1")
        valid_cv_columns = set(
            desc[0].lower() for desc in cur.description
            if desc[0].lower() not in ["id", "employee_id"]
        )
        cv_updates = {k: v for k, v in actual_updates.items() if k.lower() in valid_cv_columns}
        
        if not cv_updates:
            cur.close()
            conn.close()
            return {"success": False, "error": "Tidak ada kolom valid yang diperbarui. Periksa nama field yang diberikan."}
        
        messages = []
        
        # --- Update employee_cv table ---
        if cv_updates:
            # Check if CV record exists
            cur.execute("SELECT id FROM employee_cv WHERE employee_id = :eid", {"eid": emp_id})
            cv_row = cur.fetchone()
            
            if cv_row:
                cv_updates["emp_id"] = emp_id
                set_clause = ", ".join(f"{k} = :{k}" for k in cv_updates if k != "emp_id")
                set_clause += ", updated_at = CURRENT_TIMESTAMP"
                cur.execute(f"UPDATE employee_cv SET {set_clause} WHERE employee_id = :emp_id", cv_updates)
                messages.append(f"{len(cv_updates) - 1} field tabel employee_cv")
            else:
                # Insert minimal CV record with the given fields
                cv_updates["employee_id"] = emp_id
                cols = ", ".join(cv_updates.keys())
                placeholders = ", ".join(f":{k}" for k in cv_updates.keys())
                cur.execute(f"INSERT INTO employee_cv ({cols}) VALUES ({placeholders})", cv_updates)
                messages.append(f"{len(cv_updates) - 1} field tabel employee_cv (record baru dibuat)")
        
        conn.commit()

        # -- Re-fetch actual committed values
        updated_data = {}
        try:
            if cv_updates:
                cv_fields = [k for k in cv_updates if k not in ("emp_id", "employee_id")]
                if cv_fields:
                    cur.execute(
                        f"SELECT {', '.join(cv_fields)} FROM employee_cv WHERE employee_id = :eid",
                        {"eid": emp_id}
                    )
                    row = cur.fetchone()
                    if row:
                        for i, field in enumerate(cv_fields):
                            val = row[i]
                            updated_data[field] = val.read() if hasattr(val, "read") else val
        except Exception:
            pass  # non-critical

        cur.close()
        conn.close()

        return {
            "success": True,
            "message": f"Data resume/CV karyawan ID {emp_id} berhasil diperbarui: {', '.join(messages)}.",
            "updated_fields": updated_data 
        }
    except Exception as e:
        return {"success": False, "error": str(e), "trace": traceback.format_exc()}

# ==================== TOOL DEFINITIONS ====================

CV_TOOLS = [
    {
        "name": "update_employee_cv",
        "description": (
            "UPDATE informasi PELENGKAP (Tabel employee_cv) karyawan. "
            "Gunakan ini SAJA (bukan update_employee_by_id) untuk mengupdate/menambah detail seperti: "
            "Pendidikan, institusi, jurusan, tahun lulus, sertifikasi, keahlian, pengalaman kerja, "
            "data darurat, agama, KTP/NPWP, nomor rekening bank, serta SEMUA DETAIL POTONGAN BULANAN (BPJS, makan, transport, cicilan, dll). "
            "Data ini akan tertaut langsung ke profil CV & HR keryawan."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "emp_id": {
                    "type": "integer",
                    "description": "Database ID karyawan"
                },
                "education_level": {"type": "string", "description": "Tingkat Pendidikan (contoh: S1, SMA)"},
                "education_institution": {"type": "string", "description": "Nama Universitas / Sekolah"},
                "education_major": {"type": "string", "description": "Jurusan Utama"},
                "graduation_year": {"type": "integer", "description": "Tahun Lulus (contoh: 2022)"},
                "certifications": {"type": "string", "description": "Daftar Sertifikasi (teks bebas)"},
                "skills": {"type": "string", "description": "Daftar Keahlian/Skill (teks bebas)"},
                "work_experience": {"type": "string", "description": "Riwayat Pengalaman Kerja (teks bebas)"},
                "current_position": {"type": "string", "description": "Posisi Saat Ini di CV"},
                "current_department": {"type": "string", "description": "Departemen Saat Ini di CV"},
                "current_salary": {"type": "number", "description": "Gaji Saat Ini di CV (angka)"},
                "emergency_contact_name": {"type": "string", "description": "Nama Kontak Darurat"},
                "emergency_contact_phone": {"type": "string", "description": "No Telepon Kontak Darurat"},
                "emergency_contact_relation": {"type": "string", "description": "Hubungan Kontak Darurat (contoh: Istri)"},
                "blood_type": {"type": "string", "description": "Golongan Darah (A, B, AB, O)"},
                "religion": {"type": "string", "description": "Agama"},
                "ktp_number": {"type": "string", "description": "Nomor KTP/NIK (string)"},
                "npwp_number": {"type": "string", "description": "Nomor NPWP (string)"},
                "bank_name": {"type": "string", "description": "Nama Bank (contoh: BCA, Mandiri)"},
                "bank_account_number": {"type": "string", "description": "Nomor Rekening Bank (string)"},
                "bank_account_name": {"type": "string", "description": "Nama Pemilik Rekening Bank"},
                "deduction_bpjs_kesehatan": {"type": "number", "description": "Potongan UANG BPJS Kesehatan (angka)"},
                "deduction_bpjs_ketenagakerjaan": {"type": "number", "description": "Potongan UANG BPJS Ketenagakerjaan (angka)"},
                "deduction_meal": {"type": "number", "description": "Potongan Uang Makan (angka)"},
                "deduction_transport": {"type": "number", "description": "Potongan Uang Transport (angka)"},
                "deduction_insurance": {"type": "number", "description": "Potongan Asuransi Lain (angka)"},
                "deduction_laptop_installment": {"type": "number", "description": "Potongan Cicilan Laptop (angka)"},
                "deduction_laptop_remaining_months": {"type": "integer", "description": "Sisa Bulan Cicilan Laptop (angka)"},
                "deduction_other": {"type": "number", "description": "Potongan Lainnya (angka)"},
                "deduction_other_description": {"type": "string", "description": "Deskripsi Potongan Lainnya"},
                "total_monthly_deductions": {"type": "number", "description": "Total Potongan (angka)"},
                "notes": {"type": "string", "description": "Catatan Tambahan"}
            },
            "required": ["emp_id"]
        }
    },
    {
        "name": "get_employee_cv",
        "description": "Ambil data CV/profil LENGKAP karyawan termasuk pendidikan, sertifikasi, keahlian, pengalaman kerja, data pribadi (KTP, NPWP, golongan darah), kontak darurat, info bank, dan rincian potongan bulanan (BPJS, asuransi, cicilan laptop, dll).",
        "parameters": {
            "type": "object",
            "properties": {
                "emp_id": {
                    "type": "integer",
                    "description": "Database ID karyawan"
                }
            },
            "required": ["emp_id"]
        }
    },
    {
        "name": "analyze_employee_cv",
        "description": "ANALISA mendalam CV karyawan menggunakan AI. Mengevaluasi kekuatan, kelemahan, potensi karir, kesesuaian skill dengan posisi, dan kompensasi. Bisa difokuskan pada aspek tertentu (skills, career, compensation, performance).",
        "parameters": {
            "type": "object",
            "properties": {
                "emp_id": {
                    "type": "integer",
                    "description": "Database ID karyawan"
                },
                "focus": {
                    "type": "string",
                    "enum": ["general", "skills", "career", "compensation", "performance"],
                    "description": "Fokus analisa. Default: general.",
                    "default": "general"
                }
            },
            "required": ["emp_id"]
        }
    },
    {
        "name": "summarize_employee_cv",
        "description": "RANGKUM CV karyawan atau JAWAB PERTANYAAN spesifik tentang CV menggunakan AI. Jika ada pertanyaan, AI akan menjawab berdasarkan data CV. Jika tidak ada pertanyaan, AI akan membuat rangkuman komprehensif.",
        "parameters": {
            "type": "object",
            "properties": {
                "emp_id": {
                    "type": "integer",
                    "description": "Database ID karyawan"
                },
                "question": {
                    "type": "string",
                    "description": "Pertanyaan spesifik tentang CV. Kosongkan untuk rangkuman umum. Contoh: 'Apa sertifikasi yang dimiliki?', 'Berapa lama pengalaman kerja?'"
                }
            },
            "required": ["emp_id"]
        }
    },
    {
        "name": "manage_cv_file",
        "description": (
            "SIMPAN / KELOLA file CV karyawan ke sistem. "
            "Gunakan tool ini PERTAMA KALI saat ada file CV yang dilampirkan (attachment) oleh user, "
            "misalnya ketika user berkata: 'update karyawan dengan CV terlampir', 'upload CV ini', "
            "'ganti file CV', atau 'lampirkan CV ke profil karyawan'. "
            "action='upload' untuk file baru, action='replace' untuk mengganti file lama, action='delete' untuk menghapus. "
            "SETELAH tool ini berhasil, lanjutkan dengan extract_cv_from_file untuk mengisi data profil dari CV tersebut."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "emp_id": {
                    "type": "integer",
                    "description": "Database ID karyawan"
                },
                "action": {
                    "type": "string",
                    "enum": ["upload", "replace", "delete"],
                    "description": "Aksi: 'upload' (file baru), 'replace' (ganti file lama), 'delete' (hapus file)"
                },
                "file_path": {
                    "type": "string",
                    "description": (
                        "Path absolut ke file yang akan di-upload/replace. "
                        "Wajib untuk action upload/replace. "
                        "Ambil dari metadata attachment yang dikirim user dalam pesan (field 'path' atau 'saved_path'). "
                        "Contoh: D:\\\\PINET\\\\AI\\\\Primasistant-HR\\\\backend\\\\uploads\\\\temp\\\\namafile.pdf"
                    )
                }
            },
            "required": ["emp_id", "action"]
        }
    },
    {
        "name": "extract_cv_from_file",
        "description": (
            "BACA file CV dan UPDATE/ISI data profil karyawan secara otomatis dari isi CV tersebut. "
            "Gunakan tool ini saat user meminta: 'update informasi karyawan dari CV', 'perbarui data dari CV terlampir', "
            "'isi profil berdasarkan CV ini', 'ekstrak data CV', atau setelah manage_cv_file berhasil menyimpan file CV. "
            "AI akan membaca isi file CV (PDF/DOCX/TXT) dan mengisi kolom profil karyawan di database: "
            "pendidikan, skill, sertifikasi, pengalaman kerja, data pribadi (KTP, NPWP, golongan darah, agama), "
            "kontak darurat, dan info bank. Hanya field yang ditemukan di CV yang akan diisi. "
            "ALUR LENGKAP saat ada CV attachment: (1) manage_cv_file action=upload/replace → (2) extract_cv_from_file."
        ),
        "parameters": {
            "type": "object",
            "properties": {
                "emp_id": {
                    "type": "integer",
                    "description": "Database ID karyawan"
                },
                "file_path": {
                    "type": "string",
                    "description": (
                        "Path file CV yang akan dibaca. Opsional — jika tidak diisi, "
                        "akan otomatis menggunakan file yang sudah disimpan via manage_cv_file sebelumnya."
                    )
                }
            },
            "required": ["emp_id"]
        }
    }
]
